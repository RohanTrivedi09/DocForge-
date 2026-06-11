"""
.ipynb → .docx conversion for WordForge.
Handles markdown, syntax-coloured code, outputs, images, DataFrames, TOC, header/footer.
"""

import io
import re
import base64
import tokenize as _tokenize
import token as _token
from typing import Dict, Any, Optional, List, Tuple

import nbformat
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Sanitise ──────────────────────────────────────────────────────────────────

_ANSI_RE = re.compile(
    r"\x1b\[[0-9;]*[mGKHF]"
    r"|\x1b\[[?][0-9;]*[hl]"
    r"|\x1b\[[\x30-\x3f]*[\x20-\x2f]*[\x40-\x7e]"
    r"|\x1b."
)

def _sanitize(text: str) -> str:
    text = _ANSI_RE.sub("", text)
    return "".join(
        ch for ch in text
        if ch in ("\t", "\n", "\r") or (0x20 <= ord(ch) <= 0x7E) or ord(ch) > 0x9F
    )


# ── XML helpers ───────────────────────────────────────────────────────────────

def _set_para_shading(para, fill_hex: str):
    pPr = para._p.get_or_add_pPr()
    for ex in pPr.findall(qn("w:shd")):
        pPr.remove(ex)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex.lstrip("#"))
    pPr.append(shd)


def _add_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                border = OxmlElement(f"w:{side}")
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "4")
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), "AAAAAA")
                tcBorders.append(border)
            tcPr.append(tcBorders)


def _shade_row(row, fill_hex: str):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill_hex.lstrip("#"))
        tcPr.append(shd)


def _set_cell_shading(cell, fill_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex.lstrip("#"))
    tcPr.append(shd)


# ── Header / Footer ───────────────────────────────────────────────────────────
# A4 text width with 2.54 cm (1 inch) margins = 21 cm - 5.08 cm = 15.92 cm
# 15.92 cm × (1440 twips / 2.54 cm) = 9029 twips
_RIGHT_TWIPS  = 9029
_CENTER_TWIPS = 4514  # half of _RIGHT_TWIPS


def _reset_para(para):
    """
    Completely wipe a header/footer paragraph's XML, then return it.
    para.clear() only removes run content — this also strips pPr so no
    template tab stops or page-number fields bleed through.
    """
    p = para._p
    for child in list(p):
        p.remove(child)
    return para


def _fresh_hf_para(container):
    """Get the first paragraph of a header/footer, strip it bare, and delete the rest."""
    while len(container.paragraphs) > 1:
        p_to_remove = container.paragraphs[-1]._p
        p_to_remove.getparent().remove(p_to_remove)
    para = container.paragraphs[0]
    return _reset_para(para)


def _build_pPr_tabs(para, *tab_specs):
    """
    Create a fresh pPr with exactly ONE w:tabs element containing the
    requested tab stops.  Each spec is (val, pos_twips).
    """
    pPr = OxmlElement("w:pPr")
    tabs_el = OxmlElement("w:tabs")
    for val, pos in tab_specs:
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), val)
        tab.set(qn("w:pos"), str(pos))
        tabs_el.append(tab)
    pPr.append(tabs_el)
    para._p.append(pPr)
    return pPr


def _hf_run(para, text: str) -> Any:
    r = para.add_run(text)
    r.font.size = Pt(10)
    r.font.name = "Times New Roman"
    return r


def _page_number_field(para):
    """Inline PAGE field — three consecutive runs (begin / instr / end)."""
    for ftype, content in [("begin", None), ("instr", " PAGE "), ("end", None)]:
        run = para.add_run()
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        if ftype == "instr":
            el = OxmlElement("w:instrText")
            el.set(qn("xml:space"), "preserve")
            el.text = content
            run._r.append(el)
        else:
            el = OxmlElement("w:fldChar")
            el.set(qn("w:fldCharType"), ftype)
            run._r.append(el)


def _build_header(section, left_text: str, right_text: str):
    """Left-aligned course code | right-aligned subject, thin bottom border."""
    header = section.header
    para = _fresh_hf_para(header)

    pPr = _build_pPr_tabs(para, ("right", _RIGHT_TWIPS))

    # Bottom border on header paragraph
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)

    _hf_run(para, left_text)
    para.add_run("\t")
    _hf_run(para, right_text)


def _build_footer(section, left_text: str, right_text: str, page_numbers: bool, page_pos: str):
    """Left-aligned lab number | optional center page number | right-aligned enroll no."""
    footer = section.footer
    para = _fresh_hf_para(footer)

    if page_numbers and page_pos == "center":
        # Tab stops: center for page number, right for enrollment
        _build_pPr_tabs(para,
                        ("center", _CENTER_TWIPS),
                        ("right",  _RIGHT_TWIPS))
        _hf_run(para, left_text)
        para.add_run("\t")
        _page_number_field(para)
        para.add_run("\t")
        _hf_run(para, right_text)

    elif page_numbers and page_pos == "right":
        # Tab stops: right for page number; enrollment sits before tab
        _build_pPr_tabs(para, ("right", _RIGHT_TWIPS))
        _hf_run(para, left_text)
        para.add_run("    ")
        _hf_run(para, right_text)
        para.add_run("\t")
        _page_number_field(para)

    else:
        # No page numbers — just left tab-right
        _build_pPr_tabs(para, ("right", _RIGHT_TWIPS))
        _hf_run(para, left_text)
        para.add_run("\t")
        _hf_run(para, right_text)


# ── TOC field ─────────────────────────────────────────────────────────────────

def _insert_toc(doc: Document):
    para = doc.add_paragraph()
    run = para.add_run()
    fldBegin = OxmlElement("w:fldChar")
    fldBegin.set(qn("w:fldCharType"), "begin")
    run._r.append(fldBegin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._r.append(instr)
    fldEnd = OxmlElement("w:fldChar")
    fldEnd.set(qn("w:fldCharType"), "end")
    run._r.append(fldEnd)
    note = doc.add_paragraph()
    nr = note.add_run("(Press Ctrl+A then F9 in Word to update the Table of Contents)")
    nr.font.size = Pt(9)
    nr.italic = True
    nr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_paragraph()


# ── Inline markdown helpers ───────────────────────────────────────────────────

def _add_inline_code_run(para, text: str):
    """Courier New + light gray background shading (#E8E8E8) for inline code."""
    run = para.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(10)
    rPr = run._r.get_or_add_rPr()
    rShd = OxmlElement("w:shd")
    rShd.set(qn("w:val"), "clear")
    rShd.set(qn("w:color"), "auto")
    rShd.set(qn("w:fill"), "E8E8E8")
    rPr.append(rShd)
    return run


def _apply_inline_markdown(text: str, para):
    # Split on bold, italic, inline code, and $...$  LaTeX
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*|`.*?`|\$[^$\n]+?\$)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = para.add_run(part[2:-2]); r.bold = True
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)
        elif part.startswith("*") and part.endswith("*"):
            r = para.add_run(part[1:-1]); r.italic = True
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)
        elif part.startswith("`") and part.endswith("`"):
            _add_inline_code_run(para, part[1:-1])
        elif part.startswith("$") and part.endswith("$"):
            # LaTeX math: strip delimiters, render as italic Courier fallback
            math_text = part[1:-1]
            r = para.add_run(math_text)
            r.font.name = "Courier New"
            r.font.size = Pt(10)
            r.italic = True
        else:
            r = para.add_run(part)
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)


# ── Markdown pipe-table renderer ─────────────────────────────────────────────

def _is_md_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.endswith("|") and len(s) > 2

def _is_md_separator(line: str) -> bool:
    return bool(re.match(r"^\|[-| :]+\|$", line.strip()))

def _render_md_table(doc: Document, table_lines: List[str]):
    rows_raw = [l for l in table_lines if not _is_md_separator(l)]
    if not rows_raw:
        return
    parsed = [[c.strip() for c in row.strip("|").split("|")] for row in rows_raw]
    col_count = max(len(r) for r in parsed)
    table = doc.add_table(rows=len(parsed), cols=col_count)
    table.style = "Table Grid"
    _add_table_borders(table)
    for ri, row_data in enumerate(parsed):
        row = table.rows[ri]
        if ri == 0:
            _shade_row(row, "D9D9D9")
        elif ri % 2 == 0:
            _shade_row(row, "F7F7F7")
        for ci in range(col_count):
            cell_text = _sanitize(row_data[ci]) if ci < len(row_data) else ""
            cell = row.cells[ci]
            cell.text = cell_text
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)
                    if ri == 0:
                        run.bold = True
    doc.add_paragraph()


def _apply_heading_runs(para, text: str, level: int):
    """Format and enforce heading runs in Times New Roman, bold, and specific hex colors."""
    _apply_inline_markdown(text, para)
    HEADING_SPEC = {
        1: {"size": 16, "color": "1F3864"},
        2: {"size": 14, "color": "2E74B5"},
        3: {"size": 12, "color": "000000"},
        4: {"size": 11, "color": "000000"},
    }
    spec = HEADING_SPEC.get(level, {"size": 12, "color": "000000"})
    for run in para.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(spec["size"])
        run.bold = True
        
        # Enforce color directly in run XML to survive style inheritance
        h = spec["color"].upper()
        rPr = run._r.get_or_add_rPr()
        for existing in rPr.findall(qn("w:color")):
            rPr.remove(existing)
        color_el = OxmlElement("w:color")
        color_el.set(qn("w:val"), h)
        rPr.append(color_el)


# ── Markdown cell renderer ────────────────────────────────────────────────────

def _process_markdown(doc: Document, source: str):
    lines = source.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Heading
        m = re.match(r"^(#{1,6})\s+(.*)", line)
        if m:
            level = min(len(m.group(1)), 4)
            heading = doc.add_heading(level=level)
            heading.text = ""
            _apply_heading_runs(heading, m.group(2), level)
            i += 1; continue

        # Pipe table — collect all consecutive pipe rows
        if _is_md_table_row(line):
            table_block = []
            while i < len(lines) and (_is_md_table_row(lines[i]) or _is_md_separator(lines[i])):
                table_block.append(lines[i])
                i += 1
            _render_md_table(doc, table_block)
            continue

        # Bullet list
        if re.match(r"^[-*+]\s+", line):
            para = doc.add_paragraph(style="List Bullet")
            _apply_inline_markdown(re.sub(r"^[-*+]\s+", "", line), para)
            i += 1; continue

        # Numbered list
        if re.match(r"^\d+\.\s+", line):
            para = doc.add_paragraph(style="List Number")
            _apply_inline_markdown(re.sub(r"^\d+\.\s+", "", line), para)
            i += 1; continue

        # Display math block $$...$$
        if line.strip().startswith("$$"):
            math_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("$$"):
                math_lines.append(lines[i])
                i += 1
            i += 1  # skip closing $$
            para = doc.add_paragraph()
            r = para.add_run(" ".join(math_lines).strip())
            r.font.name = "Courier New"; r.font.size = Pt(10); r.italic = True
            continue

        if not line.strip():
            i += 1; continue

        para = doc.add_paragraph()
        _apply_inline_markdown(line, para)
        i += 1


# ── Syntax coloring ───────────────────────────────────────────────────────────

_KEYWORDS = frozenset({
    "False", "None", "True", "and", "as", "assert", "async", "await",
    "break", "class", "continue", "def", "del", "elif", "else", "except",
    "finally", "for", "from", "global", "if", "import", "in", "is",
    "lambda", "nonlocal", "not", "or", "pass", "raise", "return", "try",
    "while", "with", "yield",
})

_C_BLUE    = RGBColor(0x00, 0x00, 0xFF)   # keywords
_C_RED     = RGBColor(0xBA, 0x21, 0x21)   # strings
_C_GREEN   = RGBColor(0x3D, 0x7A, 0x00)   # comments
_C_NUM     = RGBColor(0x00, 0x80, 0x00)   # numbers
_C_BLACK   = RGBColor(0x00, 0x00, 0x00)


def _tokenize_code(source: str) -> List[Tuple[int, int, int, str, Optional[RGBColor]]]:
    """
    Returns list of (row0, col_start, col_end, text, color) tuples.
    row0 is 0-indexed line number.
    """
    segments: List[Tuple[int, int, int, str, Optional[RGBColor]]] = []
    try:
        tokens = list(_tokenize.generate_tokens(io.StringIO(source).readline))
    except _tokenize.TokenError:
        return []

    for tok in tokens:
        ttype, tstring, (srow, scol), (erow, ecol), _ = tok
        if ttype in (_token.NEWLINE, _token.NL, _token.ENDMARKER, _token.INDENT, _token.DEDENT):
            continue
        if ttype == _token.COMMENT:
            color = _C_GREEN
        elif ttype == _token.STRING:
            color = _C_RED
        elif ttype == _token.NUMBER:
            color = _C_NUM
        elif ttype == _token.NAME and tstring in _KEYWORDS:
            color = _C_BLUE
        else:
            color = _C_BLACK
        # Only handle single-line tokens for simplicity
        if srow == erow:
            segments.append((srow - 1, scol, ecol, tstring, color))
    return segments


_C_MAGIC = RGBColor(0x66, 0x66, 0x66)   # gray for magic/shell lines


def _preprocess_code_lines(source: str) -> List[Tuple[str, str]]:
    """
    Returns list of (line_text, kind) where kind is 'code', 'magic', or 'shell'.
    If a cell-magic (%%) is found, the entire source becomes magic lines.
    """
    raw_lines = source.split("\n")
    result: List[Tuple[str, str]] = []
    for line in raw_lines:
        stripped = line.strip()
        if stripped.startswith("%%"):
            return [(l, "magic") for l in raw_lines]
        elif stripped.startswith("%"):
            result.append((line, "magic"))
        elif stripped.startswith("!"):
            result.append((line, "shell"))
        else:
            result.append((line, "code"))
    return result


def _add_magic_run(para, text: str):
    r = para.add_run(text)
    r.font.name = "Courier New"; r.font.size = Pt(10)
    r.italic = True; r.font.color.rgb = _C_MAGIC


def _add_code_cell(doc: Document, source: str, execution_count: Optional[int], syntax_color: bool):
    # Label: In [N]:
    label = doc.add_paragraph()
    lr = label.add_run(f"In [{execution_count if execution_count is not None else ' '}]:")
    lr.italic = True
    lr.font.size = Pt(10)
    lr.font.name = "Times New Roman"
    lr.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    label.paragraph_format.space_after = Pt(0)

    tagged = _preprocess_code_lines(source or "")

    # Build a mapping: code-only row index → original tagged index
    code_row_to_orig: Dict[int, int] = {}
    code_row_idx = 0
    for orig_idx, (_, kind) in enumerate(tagged):
        if kind == "code":
            code_row_to_orig[code_row_idx] = orig_idx
            code_row_idx += 1

    code_only_src = "\n".join(l for l, k in tagged if k == "code")

    # Build token map keyed by original line index
    by_orig_row: Dict[int, List[Tuple[int, int, str, Optional[RGBColor]]]] = {}
    if syntax_color:
        segments = _tokenize_code(code_only_src)
        for (code_row, cs, ce, txt, col) in segments:
            orig_row = code_row_to_orig.get(code_row)
            if orig_row is not None:
                by_orig_row.setdefault(orig_row, []).append((cs, ce, txt, col))
        for row in by_orig_row:
            by_orig_row[row].sort(key=lambda x: x[0])

    for li, (raw_line, kind) in enumerate(tagged):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        _set_para_shading(para, "F2F2F2")

        if kind in ("magic", "shell"):
            _add_magic_run(para, raw_line)
            continue

        if not syntax_color:
            r = para.add_run(raw_line)
            r.font.name = "Courier New"; r.font.size = Pt(10); r.font.color.rgb = _C_BLACK
            continue

        row_segs = by_orig_row.get(li, [])
        if not row_segs:
            r = para.add_run(raw_line)
            r.font.name = "Courier New"; r.font.size = Pt(10); r.font.color.rgb = _C_BLACK
            continue

        pos = 0
        for (cs, ce, txt, color) in row_segs:
            if cs > pos:
                gap = raw_line[pos:cs]
                if gap:
                    r = para.add_run(gap)
                    r.font.name = "Courier New"; r.font.size = Pt(10); r.font.color.rgb = _C_BLACK
            r = para.add_run(txt)
            r.font.name = "Courier New"; r.font.size = Pt(10)
            if color:
                r.font.color.rgb = color
            pos = ce
        if pos < len(raw_line):
            r = para.add_run(raw_line[pos:])
            r.font.name = "Courier New"; r.font.size = Pt(10); r.font.color.rgb = _C_BLACK

    doc.add_paragraph()


# ── Text / warning output ─────────────────────────────────────────────────────

def _add_text_output(doc: Document, text: str, shaded: bool = False):
    """Render plain text output. shaded=True → yellow warning box."""
    for line in _sanitize(text).split("\n"):
        if not line.strip():
            continue
        para = doc.add_paragraph()
        run = para.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(10)
        if shaded:
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            _set_para_shading(para, "FFFBE6")
        else:
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            # Indent 0.5 cm from left
            para.paragraph_format.left_indent = Cm(0.5)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()


# ── Image output ──────────────────────────────────────────────────────────────

def _add_image_output(doc: Document, b64_data: str, caption: Optional[str] = None):
    try:
        img_bytes = base64.b64decode(b64_data)
        last = doc.add_picture(io.BytesIO(img_bytes), width=Cm(12))
        # Centre the picture paragraph
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cr = cap.add_run(caption)
            cr.font.size = Pt(9)
            cr.italic = True
            cr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        doc.add_paragraph()
    except Exception:
        doc.add_paragraph("[Image could not be embedded]")


# ── DataFrame table ───────────────────────────────────────────────────────────

def _add_dataframe_table(doc: Document, text: str):
    lines = [l for l in text.strip().split("\n") if l.strip()]
    # Filter out pandas separator rows (e.g. "---  ---")
    lines = [l for l in lines if not re.match(r"^[-\s|]+$", l)]
    if len(lines) < 2:
        _add_text_output(doc, text)
        return

    rows = [re.split(r"\s{2,}|\t", l.strip()) for l in lines]
    max_cols = max(len(r) for r in rows)
    if max_cols < 2:
        _add_text_output(doc, text)
        return

    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"
    _add_table_borders(table)

    for ri, row_data in enumerate(rows):
        row = table.rows[ri]
        if ri == 0:
            _shade_row(row, "D9D9D9")
        elif ri % 2 == 0:
            _shade_row(row, "F7F7F7")
        for ci, cell_text in enumerate(row_data):
            if ci >= max_cols:
                break
            cell = row.cells[ci]
            cell.text = _sanitize(cell_text.strip())
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    if ri == 0:
                        run.bold = True

    doc.add_paragraph()


# ── Main converter ────────────────────────────────────────────────────────────

def convert_notebook(file_bytes: bytes, cfg: Dict[str, Any]) -> bytes:
    nb = nbformat.reads(file_bytes.decode("utf-8"), as_version=4)
    doc = Document()

    # Margins
    raw_margins = cfg.get("margins", 2.5)
    margins_val = float(raw_margins.get("top", 2.5)) if isinstance(raw_margins, dict) else float(raw_margins)

    for section in doc.sections:
        section.top_margin = Cm(margins_val)
        section.bottom_margin = Cm(margins_val)
        section.left_margin = Cm(margins_val)
        section.right_margin = Cm(margins_val)

    # Header / footer
    course_code = cfg.get("courseCode", "").strip()
    subject     = cfg.get("subject", "").strip()
    lab_number  = cfg.get("labNumber", "").strip()
    enroll_no   = cfg.get("enrollNo", "").strip()
    page_numbers = bool(cfg.get("pageNumbers", False))
    page_pos     = cfg.get("pageNumberPos", "center")

    if course_code or subject or lab_number or enroll_no:
        for section in doc.sections:
            section.different_first_page_header_footer = False
            if course_code or subject:
                _build_header(section, course_code, subject)
            if lab_number or enroll_no or page_numbers:
                _build_footer(section, lab_number, enroll_no, page_numbers, page_pos)

    # Options
    syntax_color    = bool(cfg.get("syntaxColor", True))
    embed_images    = bool(cfg.get("embedImages", True))
    show_dataframes = bool(cfg.get("showDataFrames", True))
    show_warnings   = bool(cfg.get("showWarnings", False))
    add_toc         = bool(cfg.get("addToc", False))
    add_captions    = bool(cfg.get("addCaptions", True))

    # TOC
    if add_toc:
        doc.add_heading("Table of Contents", level=2)
        _insert_toc(doc)

    figure_count = 0

    # Collect doc title from first H1 markdown cell
    doc_title = ""
    for cell in nb.cells:
        if cell.get("cell_type") == "markdown":
            src = cell.get("source", "") or ""
            m = re.match(r"^#\s+(.+)", src.strip())
            if m:
                doc_title = m.group(1).strip()
                break
    if not doc_title:
        doc_title = cfg.get("filename", "notebook").replace(".docx", "").replace("_", " ").title()
    doc.core_properties.title = doc_title

    for cell in nb.cells:
        source   = cell.get("source", "") or ""
        cell_type = cell.get("cell_type", "")

        if cell_type == "markdown":
            if source.strip():
                _process_markdown(doc, source)

        elif cell_type == "raw":
            if source.strip():
                para = doc.add_paragraph()
                run = para.add_run(source)
                run.font.name = "Times New Roman"
                run.font.size = Pt(11)

        elif cell_type == "code":
            exec_count = cell.get("execution_count")

            if source.strip():
                _add_code_cell(doc, source, exec_count, syntax_color)

            for output in cell.get("outputs", []):
                otype = output.get("output_type", "")

                if otype == "stream":
                    stream_name = output.get("name", "stdout")
                    text = "".join(output.get("text", []))
                    if stream_name == "stderr":
                        if show_warnings and text.strip():
                            _add_text_output(doc, text, shaded=True)
                    else:
                        if text.strip():
                            _add_text_output(doc, text, shaded=False)

                elif otype in ("display_data", "execute_result"):
                    data = output.get("data", {})

                    embedded_image = False
                    if embed_images:
                        for mime in ("image/png", "image/jpeg", "image/gif"):
                            if mime in data:
                                figure_count += 1
                                caption = f"Figure {figure_count}" if add_captions else None
                                _add_image_output(doc, data[mime], caption)
                                embedded_image = True
                                break

                    if not embedded_image:
                        text = "".join(data.get("text/plain", []))
                        if text.strip():
                            if show_dataframes and re.search(r"\s{2,}", text) and text.count("\n") > 1:
                                _add_dataframe_table(doc, text)
                            else:
                                _add_text_output(doc, text, shaded=False)

                elif otype == "error":
                    ename  = _sanitize(output.get("ename", "Error"))
                    evalue = _sanitize(output.get("evalue", ""))
                    tb_lines = [_sanitize(l) for l in output.get("traceback", [])]
                    err_text = f"{ename}: {evalue}"
                    if tb_lines:
                        err_text += "\n" + "\n".join(tb_lines[-10:])
                    _add_text_output(doc, err_text, shaded=True)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
