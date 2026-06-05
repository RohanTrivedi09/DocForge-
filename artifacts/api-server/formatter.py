"""
python-docx formatting logic for WordForge.
Uses direct XML manipulation for heading/body colors to survive style inheritance.
"""

import io
from typing import Optional, Dict, Any

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Color helpers ─────────────────────────────────────────────────────────────

def _hex_to_rgb(hex_color: str) -> Optional[RGBColor]:
    h = hex_color.lstrip("#")
    if len(h) == 6:
        return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))
    return None


def _apply_run_color_xml(run, hex_color: str):
    """
    Apply color to a run via both the SDK and direct XML manipulation.
    Direct XML is necessary to survive style inheritance in python-docx.
    """
    h = hex_color.lstrip("#").upper()
    rgb = _hex_to_rgb(hex_color)
    if rgb:
        run.font.color.rgb = rgb

    rPr = run._r.get_or_add_rPr()
    # Remove existing color element if present to avoid duplicates
    for existing in rPr.findall(qn("w:color")):
        rPr.remove(existing)
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), h)
    rPr.append(color_el)


def _apply_heading_style(paragraph, size_pt: float, hex_color: str, font_family: str, bold: bool = True):
    """Apply size, color, and font to all runs in a heading paragraph via XML."""
    for run in paragraph.runs:
        run.font.name = font_family
        run.font.size = Pt(size_pt)
        run.bold = bold
        _apply_run_color_xml(run, hex_color)


# ── Page number field helpers ─────────────────────────────────────────────────

def _add_field(run, field_code: str):
    """Insert a Word field (e.g. PAGE, NUMPAGES) into a run."""
    fldBegin = OxmlElement("w:fldChar")
    fldBegin.set(qn("w:fldCharType"), "begin")
    run._r.append(fldBegin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    run._r.append(instr)

    fldEnd = OxmlElement("w:fldChar")
    fldEnd.set(qn("w:fldCharType"), "end")
    run._r.append(fldEnd)


def _set_header_footer_text(hf, text: str, font_size: int = 10):
    """Set plain text in a header or footer paragraph."""
    hf.paragraphs[0].clear()
    run = hf.paragraphs[0].add_run(text)
    run.font.size = Pt(font_size)


def _add_page_number(footer, position: str = "center", page_x_of_y: bool = False):
    """Add automatic page number field (or 'Page X of Y') to a footer paragraph."""
    para = footer.paragraphs[0]
    para.clear()

    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }
    para.alignment = align_map.get(position, WD_ALIGN_PARAGRAPH.CENTER)

    if page_x_of_y:
        run1 = para.add_run("Page ")
        run1.font.size = Pt(10)
        run2 = para.add_run()
        run2.font.size = Pt(10)
        _add_field(run2, "PAGE")
        run3 = para.add_run(" of ")
        run3.font.size = Pt(10)
        run4 = para.add_run()
        run4.font.size = Pt(10)
        _add_field(run4, "NUMPAGES")
    else:
        run = para.add_run()
        run.font.size = Pt(10)
        _add_field(run, "PAGE")


LINE_SPACING_MAP = {
    "1.0": 1.0,
    "1.15": 1.15,
    "1.5": 1.5,
    "2.0": 2.0,
}

HEADING_STYLE_MAP = {
    "h1": "Heading 1",
    "h2": "Heading 2",
    "h3": "Heading 3",
    "h4": "Heading 4",
}


# ── Core formatter ────────────────────────────────────────────────────────────

def apply_formatting(file_bytes: Optional[bytes], cfg: Dict[str, Any]) -> bytes:
    """
    Apply formatting settings to a .docx.
    If file_bytes is None, creates a sample document.
    """
    if file_bytes:
        doc = Document(io.BytesIO(file_bytes))
    else:
        doc = _create_sample_document()

    font_family = cfg.get("fontFamily", "Times New Roman")
    body_size = float(cfg.get("bodySize", 12))
    body_color = cfg.get("bodyColor", "#000000")
    line_spacing = str(cfg.get("lineSpacing", "1.5"))
    para_before = float(cfg.get("spacingBefore", 0))
    para_after = float(cfg.get("spacingAfter", 6))

    heading_sizes = {
        "h1": float(cfg.get("h1Size", 16)),
        "h2": float(cfg.get("h2Size", 14)),
        "h3": float(cfg.get("h3Size", 12)),
        "h4": float(cfg.get("h4Size", 11)),
    }
    heading_colors = {
        "h1": cfg.get("h1Color", "#000000"),
        "h2": cfg.get("h2Color", "#000000"),
        "h3": cfg.get("h3Color", "#000000"),
        "h4": cfg.get("h4Color", "#000000"),
    }

    margins = cfg.get("margins", {
        "top": cfg.get("marginTop", 2.5),
        "bottom": cfg.get("marginBottom", 2.5),
        "left": cfg.get("marginLeft", 2.5),
        "right": cfg.get("marginRight", 2.5),
    })

    header_cfg = cfg.get("header", {})
    footer_cfg = cfg.get("footer", {})

    # Apply margins and header/footer to each section
    for section in doc.sections:
        section.top_margin = Cm(float(margins.get("top", cfg.get("marginTop", 2.5))))
        section.bottom_margin = Cm(float(margins.get("bottom", cfg.get("marginBottom", 2.5))))
        section.left_margin = Cm(float(margins.get("left", cfg.get("marginLeft", 2.5))))
        section.right_margin = Cm(float(margins.get("right", cfg.get("marginRight", 2.5))))

        diff_first = bool(cfg.get("diffFirstPage") or header_cfg.get("differentFirstPage"))
        diff_odd_even = bool(cfg.get("diffOddEven") or header_cfg.get("differentOddEven"))

        section.different_first_page_header_footer = diff_first
        section.different_odd_even_pages = diff_odd_even

        header_text = cfg.get("headerText") or header_cfg.get("text", "")
        footer_text = cfg.get("footerText") or footer_cfg.get("text", "")
        page_numbers = cfg.get("pageNumbers", footer_cfg.get("pageNumbers", True))
        page_num_pos = cfg.get("pageNumberPos", footer_cfg.get("pageNumberPosition", "center"))
        page_x_of_y = bool(cfg.get("pageXofY", False))

        if header_text:
            _set_header_footer_text(section.header, header_text)
            if diff_odd_even:
                _set_header_footer_text(section.even_page_header, header_text)

        if footer_text:
            _set_header_footer_text(section.footer, footer_text)

        if page_numbers:
            _add_page_number(section.footer, page_num_pos, page_x_of_y)
            if diff_odd_even:
                _add_page_number(section.even_page_footer, page_num_pos, page_x_of_y)

    spacing_val = LINE_SPACING_MAP.get(line_spacing, 1.5)

    # Apply paragraph styles
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""

        applied_heading = False
        for key, style in HEADING_STYLE_MAP.items():
            if style_name == style:
                applied_heading = True
                size = heading_sizes[key]
                color = heading_colors[key]
                bold = key in ("h3", "h4")  # H3/H4 are bold by default in the spec

                # Ensure there's at least one run
                if not para.runs:
                    para.add_run(para.text)

                _apply_heading_style(para, size, color, font_family, bold=True)
                break

        if not applied_heading:
            for run in para.runs:
                run.font.name = font_family
                run.font.size = Pt(body_size)
                _apply_run_color_xml(run, body_color)

        pf = para.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = spacing_val
        pf.space_before = Pt(para_before)
        pf.space_after = Pt(para_after)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _create_sample_document() -> Document:
    """Create a sample document with placeholder content."""
    doc = Document()
    doc.add_heading("Document Title", level=1)
    doc.add_heading("1. Section Heading", level=2)
    doc.add_paragraph(
        "This is sample body text. WordForge will apply your chosen formatting "
        "to the entire document — font, color, spacing, margins, and heading hierarchy."
    )
    doc.add_heading("1.1 Subsection", level=3)
    doc.add_paragraph(
        "Another paragraph of body text. Replace this with your own content, "
        "or upload an existing .docx file using the button above."
    )
    doc.add_heading("1.1.1 Sub-subsection", level=4)
    doc.add_paragraph(
        "Formatting rules are applied consistently throughout every paragraph "
        "and heading in the document."
    )
    return doc
