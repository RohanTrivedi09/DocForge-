"""
IEEE Conference Paper formatter for WordForge.
Converts a standard .docx into IEEE two-column format.
"""

import io
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_columns(section, num_cols: int = 2, spacing_cm: float = 0.63):
    """Set the number of columns on a section via XML."""
    sectPr = section._sectPr
    for existing in sectPr.findall(qn('w:cols')):
        sectPr.remove(existing)
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), str(num_cols))
    cols.set(qn('w:space'), str(int(spacing_cm * 567)))
    cols.set(qn('w:equalWidth'), '1')
    sectPr.append(cols)


def format_ieee(file_bytes: Optional[bytes], cfg: dict = None) -> bytes:
    """
    Apply IEEE conference paper formatting to a .docx document.
    """
    cfg = cfg or {}

    if file_bytes:
        doc = Document(io.BytesIO(file_bytes))
    else:
        doc = _create_ieee_sample()

    for section in doc.sections:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Cm(1.9)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(1.78)
        section.right_margin = Cm(1.78)
        section.header_distance = Cm(1.25)
        section.footer_distance = Cm(1.25)
        _set_columns(section, num_cols=2, spacing_cm=0.63)

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else 'Normal'
        pf = para.paragraph_format

        if style_name in ('Title',):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.space_before = Pt(0)
            pf.space_after = Pt(12)
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(24)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
            continue

        if style_name == 'Heading 1':
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.space_before = Pt(12)
            pf.space_after = Pt(6)
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
            pf.keep_with_next = True
            for run in para.runs:
                run.text = run.text.upper()
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
                rPr = run._r.get_or_add_rPr()
                smallCaps = OxmlElement('w:smallCaps')
                smallCaps.set(qn('w:val'), 'true')
                rPr.append(smallCaps)
            continue

        if style_name == 'Heading 2':
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.space_before = Pt(10)
            pf.space_after = Pt(4)
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
            pf.keep_with_next = True
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                run.font.bold = True
                run.font.italic = True
                run.font.color.rgb = RGBColor(0, 0, 0)
            continue

        if style_name == 'Heading 3':
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.space_before = Pt(8)
            pf.space_after = Pt(4)
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
            pf.keep_with_next = True
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                run.font.italic = True
                run.font.color.rgb = RGBColor(0, 0, 0)
            continue

        if style_name == 'Heading 4':
            pf.space_before = Pt(6)
            pf.space_after = Pt(4)
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                run.font.italic = True
                run.font.color.rgb = RGBColor(0, 0, 0)
            continue

        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.first_line_indent = Cm(0.35)

        for run in para.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 0, 0)

    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        if footer.paragraphs:
            p = footer.paragraphs[0]
            p.clear()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
            fldBegin = OxmlElement('w:fldChar')
            fldBegin.set(qn('w:fldCharType'), 'begin')
            run._r.append(fldBegin)
            instr = OxmlElement('w:instrText')
            instr.set(qn('xml:space'), 'preserve')
            instr.text = 'PAGE'
            run._r.append(instr)
            fldEnd = OxmlElement('w:fldChar')
            fldEnd.set(qn('w:fldCharType'), 'end')
            run._r.append(fldEnd)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _create_ieee_sample() -> Document:
    """Create a sample IEEE-formatted document."""
    doc = Document()
    doc.add_paragraph('Sample IEEE Conference Paper Title', style='Title')
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph(
        'This document provides a sample template for IEEE conference papers. '
        'Upload your own .docx file to convert it to IEEE two-column format.'
    )
    doc.add_heading('I. Introduction', level=1)
    doc.add_paragraph(
        'This is the introduction section. IEEE papers use a two-column format '
        'with Times New Roman 10pt font and single line spacing.'
    )
    doc.add_heading('A. Background', level=2)
    doc.add_paragraph(
        'Subsection content goes here. Headings at this level are italic and bold.'
    )
    doc.add_heading('II. Methodology', level=1)
    doc.add_paragraph(
        'Description of your research methodology.'
    )
    doc.add_heading('III. Results', level=1)
    doc.add_paragraph('Present your findings here.')
    doc.add_heading('IV. Conclusion', level=1)
    doc.add_paragraph('Summarize your conclusions.')
    doc.add_heading('References', level=1)
    doc.add_paragraph('[1] A. Author, "Paper title," Journal, vol. 1, pp. 1-10, 2024.')
    doc.add_paragraph('[2] B. Author, "Another paper," Proc. Conf., 2024, pp. 1-5.')
    return doc
