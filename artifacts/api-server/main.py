import os
import io
import json
import zipfile
import traceback
from pathlib import Path

from dotenv import load_dotenv
load_dotenv(dotenv_path=Path(__file__).resolve().parent.parent.parent / ".env")

from fastapi import FastAPI, File, UploadFile, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import Optional, List

from formatter import apply_formatting
from nb_converter import convert_notebook
from ai_hints import get_suggestion
from cover_page import add_cover_page
from sop_formatter import generate_sop_document
from docx import Document
from ieee_formatter import format_ieee

app = FastAPI(title="WordForge API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)


# ── Health check ──────────────────────────────────────────────────────────────

@app.get("/api/healthz")
async def healthz():
    return {"status": "ok"}


# ── Format document ──────────────────────────────────────────────────────────

@app.post("/api/format-doc")
async def format_doc(
    settings: str = Form(...),
    file: Optional[UploadFile] = File(None),
):
    try:
        cfg = json.loads(settings)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid settings JSON")

    try:
        file_bytes = await file.read() if file else None
        result_bytes = apply_formatting(file_bytes, cfg)

        # Optionally prepend a cover page
        cover_cfg = cfg.get("coverPage")
        if cover_cfg and cover_cfg.get("enabled"):
            doc = Document(io.BytesIO(result_bytes))
            add_cover_page(doc, cover_cfg)
            buf = io.BytesIO()
            doc.save(buf)
            result_bytes = buf.getvalue()

        filename = cfg.get("filename", "submission.docx")
        if not filename.endswith(".docx"):
            filename += ".docx"

        return StreamingResponse(
            io.BytesIO(result_bytes),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


# ── Convert notebook(s) ───────────────────────────────────────────────────────

@app.post("/api/convert-notebook")
async def convert_nb(
    settings: str = Form("{}"),
    files: List[UploadFile] = File(...),
):
    try:
        cfg = json.loads(settings)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid settings JSON")

    if not files:
        raise HTTPException(status_code=400, detail="No files uploaded")

    if len(files) > 5:
        raise HTTPException(status_code=400, detail="Maximum 5 notebooks per batch")

    try:
        if len(files) == 1:
            # Single file — return docx directly
            f = files[0]
            file_bytes = await f.read()
            result = convert_notebook(file_bytes, cfg)
            stem = Path(f.filename or "notebook").stem
            filename = cfg.get("filename") or f"{stem}.docx"
            if not filename.endswith(".docx"):
                filename += ".docx"
            return StreamingResponse(
                io.BytesIO(result),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f'attachment; filename="{filename}"'},
            )
        else:
            # Batch — return zip
            zip_buf = io.BytesIO()
            with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
                for f in files:
                    file_bytes = await f.read()
                    result = convert_notebook(file_bytes, cfg)
                    stem = Path(f.filename or "notebook").stem
                    zf.writestr(f"{stem}.docx", result)
            zip_buf.seek(0)
            return StreamingResponse(
                zip_buf,
                media_type="application/zip",
                headers={"Content-Disposition": 'attachment; filename="notebooks.zip"'},
            )
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


# ── Parse Document for Live Preview ───────────────────────────────────────────

@app.post("/api/parse-doc")
async def parse_doc_for_preview(file: UploadFile = File(...)):
    try:
        file_bytes = await file.read()
        doc = Document(io.BytesIO(file_bytes))
        elements = []

        from docx.oxml.table import CT_Tbl
        from docx.oxml.text.paragraph import CT_P
        from docx.table import Table
        from docx.text.paragraph import Paragraph
        from docx.oxml.ns import qn

        def process_element(element):
            for child in element.iterchildren():
                if isinstance(child, CT_P):
                    para = Paragraph(child, doc)
                    text = para.text.strip()
                    if not text:
                        continue
                    style = para.style.name if para.style else "Normal"
                    
                    el_type = "paragraph"
                    level = 0
                    if style.startswith("Heading 1"):
                        el_type = "h1"
                        level = 1
                    elif style.startswith("Heading 2"):
                        el_type = "h2"
                        level = 2
                    elif style.startswith("Heading 3"):
                        el_type = "h3"
                        level = 3
                    elif style.startswith("Heading 4"):
                        el_type = "h4"
                        level = 4
                    elif style.startswith("List Bullet") or style.startswith("List"):
                        el_type = "bullet"
                    
                    elements.append({
                        "type": el_type,
                        "level": level,
                        "text": text
                    })
                elif isinstance(child, CT_Tbl):
                    table = Table(child, doc)
                    rows_data = []
                    for row in table.rows:
                        row_cells = [cell.text.strip() for cell in row.cells]
                        rows_data.append(row_cells)
                    if rows_data:
                        elements.append({
                            "type": "table",
                            "headers": rows_data[0],
                            "rows": rows_data[1:]
                        })
                elif child.tag.endswith('sdt'):
                    sdtContent = child.find(qn('w:sdtContent'))
                    if sdtContent is not None:
                        process_element(sdtContent)

        process_element(doc.element.body)
        return {"elements": elements}
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


# ── Format IEEE Conference Paper ──────────────────────────────────────────────

@app.post("/api/format-ieee")
async def format_ieee_paper(
    settings: str = Form("{}"),
    file: Optional[UploadFile] = File(None),
):
    try:
        cfg = json.loads(settings)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid settings JSON")
    try:
        file_bytes = await file.read() if file else None
        result_bytes = format_ieee(file_bytes, cfg)
        filename = cfg.get("filename", "ieee_paper.docx")
        if not filename.endswith(".docx"):
            filename += ".docx"
        return StreamingResponse(
            io.BytesIO(result_bytes),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


# ── Format SOP (Gujarat Govt.) ────────────────────────────────────────────────

@app.post("/api/format-sop")
async def format_sop(
    chapter_number: str = Form("1"),
    chapter_title: str = Form(""),
    content: str = Form(""),
    reference_docx: Optional[UploadFile] = File(None),
):
    try:
        ref_bytes = await reference_docx.read() if reference_docx else None
        docx_bytes, report = generate_sop_document(
            content=content,
            chapter_number=chapter_number,
            chapter_title=chapter_title,
            reference_docx_bytes=ref_bytes,
        )
        import json as _json
        filename = f"SOP_Chapter_{chapter_number}.docx"
        return StreamingResponse(
            io.BytesIO(docx_bytes),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{filename}"',
                "X-SOP-Validation": _json.dumps(report),
                "Access-Control-Expose-Headers": "X-SOP-Validation",
            },
        )
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


# ── AI suggestions ────────────────────────────────────────────────────────────

class AiSuggestInput(BaseModel):
    trigger: str
    context: Optional[str] = None


@app.post("/api/ai-suggest")
async def ai_suggest(body: AiSuggestInput):
    try:
        suggestion = get_suggestion(body.trigger, body.context or "")
        return {"suggestion": suggestion}
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


# ── AI Document Analysis ──────────────────────────────────────────────────────

@app.post("/api/ai-analyze")
async def ai_analyze_doc(
    file: UploadFile = File(...),
    settings: str = Form("{}"),
):
    """Analyze an uploaded document and return specific, actionable tips."""
    try:
        file_bytes = await file.read()
        doc = Document(io.BytesIO(file_bytes))
        stats = {
            "total_paragraphs": 0,
            "headings": {"h1": 0, "h2": 0, "h3": 0, "h4": 0},
            "fonts_used": set(),
            "font_sizes_used": set(),
            "has_tables": False,
            "table_count": 0,
            "word_count": 0,
            "empty_paragraphs": 0,
            "styles_used": set(),
            "inconsistent_fonts": False,
            "has_headers": False,
            "has_footers": False,
            "heading_texts": [],
        }
        body_fonts = []
        for para in doc.paragraphs:
            stats["total_paragraphs"] += 1
            text = para.text.strip()
            if not text:
                stats["empty_paragraphs"] += 1
                continue
            stats["word_count"] += len(text.split())
            style_name = para.style.name if para.style else "Normal"
            stats["styles_used"].add(style_name)
            if style_name.startswith("Heading 1"):
                stats["headings"]["h1"] += 1
                stats["heading_texts"].append(f"H1: {text[:80]}")
            elif style_name.startswith("Heading 2"):
                stats["headings"]["h2"] += 1
                stats["heading_texts"].append(f"H2: {text[:80]}")
            elif style_name.startswith("Heading 3"):
                stats["headings"]["h3"] += 1
            elif style_name.startswith("Heading 4"):
                stats["headings"]["h4"] += 1
            for run in para.runs:
                if run.font.name:
                    stats["fonts_used"].add(run.font.name)
                    if not style_name.startswith("Heading"):
                        body_fonts.append(run.font.name)
                if run.font.size:
                    try:
                        size_pt = run.font.size.pt
                    except Exception:
                        size_pt = float(run.font.size) / 12700
                    stats["font_sizes_used"].add(round(size_pt, 1))
        from docx.oxml.table import CT_Tbl
        for child in doc.element.body.iterchildren():
            if isinstance(child, CT_Tbl):
                stats["has_tables"] = True
                stats["table_count"] += 1
        for section in doc.sections:
            if section.header and section.header.paragraphs:
                for p in section.header.paragraphs:
                    if p.text.strip():
                        stats["has_headers"] = True
            if section.footer and section.footer.paragraphs:
                for p in section.footer.paragraphs:
                    if p.text.strip():
                        stats["has_footers"] = True
        if len(set(body_fonts)) > 1:
            stats["inconsistent_fonts"] = True
        page_est = max(1, stats["word_count"] // 250)
        stats["fonts_used"] = list(stats["fonts_used"])
        stats["font_sizes_used"] = sorted(list(stats["font_sizes_used"]))
        stats["styles_used"] = list(stats["styles_used"])
        try:
            cfg = json.loads(settings)
        except Exception:
            cfg = {}
        from ai_hints import _get_client
        client = _get_client()
        prompt = f"""You are a document formatting expert analyzing a student's Word document for academic submission.

Document statistics:
- Paragraphs: {stats['total_paragraphs']}, Words: {stats['word_count']}, ~{page_est} pages
- Empty paragraphs: {stats['empty_paragraphs']}
- H1: {stats['headings']['h1']}, H2: {stats['headings']['h2']}, H3: {stats['headings']['h3']}, H4: {stats['headings']['h4']}
- Headings: {'; '.join(stats['heading_texts'][:10])}
- Fonts: {', '.join(stats['fonts_used']) or 'None detected'}
- Sizes: {', '.join(str(s) + 'pt' for s in stats['font_sizes_used']) or 'None detected'}
- Inconsistent body fonts: {stats['inconsistent_fonts']}
- Tables: {stats['table_count']}, Headers: {stats['has_headers']}, Footers: {stats['has_footers']}

Current settings: Font={cfg.get('fontFamily','?')}, Size={cfg.get('bodySize','?')}pt, Spacing={cfg.get('lineSpacing','?')}

Provide exactly 5 actionable tips. Reference specific data from the document. Format as JSON array with keys "title" (3-5 words), "tip" (1-2 sentences), "severity" ("error"/"warning"/"info"). Return ONLY the JSON array."""
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=prompt,
        )
        result_text = (response.text or "").strip()
        try:
            if result_text.startswith('```'):
                result_text = result_text.split('\n', 1)[1].rsplit('```', 1)[0].strip()
            tips = json.loads(result_text)
        except Exception:
            tips = [{"title": "Analysis", "tip": result_text[:500], "severity": "info"}]
        return {"stats": stats, "tips": tips}
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


if __name__ == "__main__":
    import uvicorn
    port = int(os.environ.get("PORT", "8080"))
    uvicorn.run(app, host="0.0.0.0", port=port)
