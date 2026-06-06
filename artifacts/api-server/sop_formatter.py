"""
Gujarat Government SOP document generator for WordForge.
All color/size values are spec-locked and must not be approximated.
"""

import io
import re
import zipfile
from typing import Optional, List, Tuple, Dict

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Spec constants ─────────────────────────────────────────────────────────────

HEADING_SPEC: Dict[int, dict] = {
    1: {"size": 14, "color": "000080", "bold": True},
    2: {"size": 14, "color": "C00000", "bold": True},
    3: {"size": 12, "color": "1F4D78", "bold": True},
    4: {"size": 12, "color": "1F4D78", "bold": True},
}

ANNEXURE_SPEC = {"size": 14, "color": "000080", "bold": True}

_BODY_FONT = "Times New Roman"
_BODY_SIZE = 12
_BODY_COLOR = RGBColor(0x00, 0x00, 0x00)


# ── XML helpers ────────────────────────────────────────────────────────────────

def _force_color(run, hex6: str):
    """Apply color to run via XML — survives all style inheritance."""
    h = hex6.lstrip("#").upper()
    try:
        run.font.color.rgb = RGBColor(
            int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
        )
    except Exception:
        pass
    rPr = run._r.get_or_add_rPr()
    for existing in rPr.findall(qn("w:color")):
        rPr.remove(existing)
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), h)
    rPr.append(color_el)


def _set_cell_shading(cell, fill_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex.lstrip("#").upper())
    tc_pr.append(shd)


def _set_cell_padding(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side in ["top", "bottom"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), "80")
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    for side in ["left", "right"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), "120")
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tc_pr.append(mar)


# ── Body / heading formatters ──────────────────────────────────────────────────

def _apply_body_fmt(para):
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = para.paragraph_format
    pf.line_spacing = Pt(_BODY_SIZE * 1.5)
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)
    pf.first_line_indent = Pt(0)
    for run in para.runs:
        run.font.name = _BODY_FONT
        run.font.size = Pt(_BODY_SIZE)
        _force_color(run, "000000")


def _apply_heading_sop(para, level: int):
    spec = HEADING_SPEC.get(level, HEADING_SPEC[3])
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = para.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    for run in para.runs:
        run.font.name = _BODY_FONT
        run.font.size = Pt(spec["size"])
        run.font.bold = spec["bold"]
        _force_color(run, spec["color"])


# ── Paragraph builders ─────────────────────────────────────────────────────────

def _add_body_para(doc: Document, text: str):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = para.paragraph_format
    pf.line_spacing = Pt(_BODY_SIZE * 1.5)
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)
    run = para.add_run(text)
    run.font.name = _BODY_FONT
    run.font.size = Pt(_BODY_SIZE)
    _force_color(run, "000000")
    return para


def _add_heading_para(doc: Document, text: str, level: int):
    spec = HEADING_SPEC.get(level, HEADING_SPEC[3])
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = para.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    run = para.add_run(text)
    run.font.name = _BODY_FONT
    run.font.size = Pt(spec["size"])
    run.font.bold = spec["bold"]
    _force_color(run, spec["color"])
    return para


def _add_sop_bullet(doc: Document, text: str):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = para.paragraph_format
    pf.left_indent = Inches(0.4)
    pf.first_line_indent = Inches(-0.2)
    pf.space_after = Pt(6)
    pf.line_spacing = Pt(18)
    run = para.add_run(f"\u2192 {text}")
    run.font.name = _BODY_FONT
    run.font.size = Pt(_BODY_SIZE)
    _force_color(run, "000000")
    return para


def _add_objective_para(doc: Document, label: str, body_text: str):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = para.paragraph_format
    pf.line_spacing = Pt(_BODY_SIZE * 1.5)
    pf.space_after = Pt(6)
    label_run = para.add_run(label + " ")
    label_run.font.name = _BODY_FONT
    label_run.font.size = Pt(_BODY_SIZE)
    label_run.font.bold = True
    label_run.font.italic = True
    _force_color(label_run, "000080")
    body_run = para.add_run(body_text)
    body_run.font.name = _BODY_FONT
    body_run.font.size = Pt(_BODY_SIZE)
    body_run.font.bold = False
    body_run.font.italic = False
    _force_color(body_run, "000000")
    return para


def _add_reference_para(doc: Document, body_text: str):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = para.paragraph_format
    pf.line_spacing = Pt(_BODY_SIZE * 1.5)
    pf.space_after = Pt(6)
    label_run = para.add_run("Reference: ")
    label_run.font.name = _BODY_FONT
    label_run.font.size = Pt(_BODY_SIZE)
    label_run.font.bold = True
    _force_color(label_run, "000080")
    body_run = para.add_run(body_text)
    body_run.font.name = _BODY_FONT
    body_run.font.size = Pt(_BODY_SIZE)
    body_run.font.bold = False
    _force_color(body_run, "000000")
    return para


def _add_citation(doc: Document, text: str):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = para.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    run = para.add_run(text)
    run.font.name = _BODY_FONT
    run.font.size = Pt(10)
    run.font.italic = True
    _force_color(run, "404040")
    return para


def _add_sop_table(doc: Document, headers: List[str], rows: List[List[str]]):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "9026")
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = ""
        _set_cell_shading(cell, "2C5282")
        _set_cell_padding(cell)
        run = cell.paragraphs[0].add_run(h)
        run.font.name = _BODY_FONT
        run.font.size = Pt(10)
        run.font.bold = True
        _force_color(run, "FFFFFF")

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ""
            _set_cell_shading(cell, "F7FAFC")
            _set_cell_padding(cell)
            run = cell.paragraphs[0].add_run(str(val))
            run.font.name = _BODY_FONT
            run.font.size = Pt(10)
            _force_color(run, "000000")


# ── Cover page ────────────────────────────────────────────────────────────────

def _generate_sop_cover(doc: Document, chapter_number: str, chapter_title: str):
    # Blank space (~40% down the page)
    for _ in range(8):
        doc.add_paragraph()

    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(f"Chapter {chapter_number}")
    r1.font.name = _BODY_FONT; r1.font.size = Pt(24); r1.font.bold = True
    _force_color(r1, "1A365D")

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(chapter_title)
    r2.font.name = _BODY_FONT; r2.font.size = Pt(18); r2.font.bold = True
    _force_color(r2, "002060")

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Standard Operating Procedure")
    r3.font.name = _BODY_FONT; r3.font.size = Pt(14); r3.font.italic = True
    _force_color(r3, "002060")

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run("Gujarat State Police Wireless Grid")
    r4.font.name = _BODY_FONT; r4.font.size = Pt(12)
    _force_color(r4, "1A365D")

    doc.add_page_break()

    # First page has no header
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    section.first_page_header.paragraphs[0].clear()


# ── Running header ─────────────────────────────────────────────────────────────

def _set_sop_running_header(section, chapter_title: str):
    header = section.header
    para = header.paragraphs[0]
    para.clear()

    run = para.add_run(chapter_title.upper())
    run.font.name = _BODY_FONT
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.italic = True
    _force_color(run, "000080")

    # Navy bottom border
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000080")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── Reference docx color extractor ────────────────────────────────────────────

def extract_and_verify_colors(docx_bytes: bytes) -> Dict[str, int]:
    """Unpack docx XML and count all hex color values found."""
    found: Dict[str, int] = {}
    try:
        with zipfile.ZipFile(io.BytesIO(docx_bytes)) as z:
            for name in z.namelist():
                if name.endswith(".xml"):
                    content = z.read(name).decode("utf-8", errors="ignore")
                    colors = re.findall(r'w:val="([0-9A-Fa-f]{6})"', content)
                    for c in colors:
                        cu = c.upper()
                        found[cu] = found.get(cu, 0) + 1
    except Exception:
        pass
    return found


# ── Content parser ─────────────────────────────────────────────────────────────

def _parse_content(doc: Document, content: str):
    """
    Parse pasted chapter text and render with SOP styles.

    Supported patterns:
      # Title            → H1
      ## Title           → H2
      ### Title          → H3
      #### Title         → H4
      -> text            → bullet
      - text / * text    → bullet
      Objective: ...     → objective para (bold italic navy label)
      Reference: ...     → reference para (bold navy label)
      Citation: ...      → 10pt italic gray
      | col | col |      → table (pipe-delimited)
      (blank line)       → spacer (skipped)
      (other)            → body text
    """
    lines = content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,4})\s+(.*)", stripped)
        if m:
            level = min(len(m.group(1)), 4)
            _add_heading_para(doc, m.group(2), level)
            i += 1
            continue

        # Pipe table — collect block
        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            # Parse: skip separator rows
            rows_raw = [l for l in table_lines if not re.match(r"^\|[-| :]+\|$", l)]
            if rows_raw:
                parsed = [[c.strip() for c in r.strip("|").split("|")] for r in rows_raw]
                if len(parsed) >= 2:
                    _add_sop_table(doc, parsed[0], parsed[1:])
                elif len(parsed) == 1:
                    _add_sop_table(doc, parsed[0], [])
            continue

        # Bullets
        if re.match(r"^(->\s+|-\s+|\*\s+)", stripped):
            text = re.sub(r"^(->\s+|-\s+|\*\s+)", "", stripped)
            _add_sop_bullet(doc, text)
            i += 1
            continue

        # Objective / Reference / Citation special labels
        obj_m = re.match(r"^(Objective\s*:)\s*(.*)", stripped, re.IGNORECASE)
        if obj_m:
            _add_objective_para(doc, obj_m.group(1), obj_m.group(2))
            i += 1
            continue

        ref_m = re.match(r"^Reference\s*:\s*(.*)", stripped, re.IGNORECASE)
        if ref_m:
            _add_reference_para(doc, ref_m.group(1))
            i += 1
            continue

        cit_m = re.match(r"^Citation\s*:\s*(.*)", stripped, re.IGNORECASE)
        if cit_m:
            _add_citation(doc, cit_m.group(1))
            i += 1
            continue

        # Default — body text
        _add_body_para(doc, stripped)
        i += 1


# ── Post-generation validator ──────────────────────────────────────────────────

def validate_sop_output(doc: Document) -> dict:
    issues = []
    para_count = len(doc.paragraphs)
    h1_count = sum(1 for p in doc.paragraphs if p.style.name == "Heading 1")
    h2_count = sum(1 for p in doc.paragraphs if p.style.name == "Heading 2")

    expected_colors = {
        "Heading 1": "000080",
        "Heading 2": "C00000",
        "Heading 3": "1F4D78",
    }
    for p in doc.paragraphs:
        if p.style.name in expected_colors and p.runs:
            try:
                color = str(p.runs[0].font.color.rgb).upper()
                exp = expected_colors[p.style.name]
                if color != exp.upper():
                    issues.append(
                        f"{p.style.name} color mismatch: got {color}, expected {exp}"
                    )
            except Exception:
                pass

    return {
        "paragraph_count": para_count,
        "h1_count": h1_count,
        "h2_count": h2_count,
        "issues": issues,
    }


# ── Main entry point ───────────────────────────────────────────────────────────

def generate_sop_document(
    content: str,
    chapter_number: str,
    chapter_title: str,
    reference_docx_bytes: Optional[bytes] = None,
) -> Tuple[bytes, dict]:
    """
    Generate a Gujarat Govt. SOP .docx from pasted chapter content.
    Returns (docx_bytes, validation_report).
    """
    doc = Document()

    # Page setup — A4 with 1-inch margins
    for section in doc.sections:
        section.page_width  = Inches(8.27)
        section.page_height = Inches(11.69)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)

    # Reference docx color verification (informational — logged in report)
    ref_colors: Dict[str, int] = {}
    if reference_docx_bytes:
        ref_colors = extract_and_verify_colors(reference_docx_bytes)

    # Cover page + running header
    _generate_sop_cover(doc, chapter_number, chapter_title)
    _set_sop_running_header(doc.sections[0], chapter_title)

    # Parse and render content
    if content.strip():
        _parse_content(doc, content)

    # Validate output
    report = validate_sop_output(doc)
    if ref_colors:
        report["reference_colors_found"] = ref_colors

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue(), report
